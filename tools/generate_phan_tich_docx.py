# -*- coding: utf-8 -*-
"""One-off: táº¡o PHAN_TICH_LOCAL_NAV_BOT.docx â€” cháº¡y: python tools/generate_phan_tich_docx.py"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent.parent / "PHAN_TICH_LOCAL_NAV_BOT.docx"


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def main():
    doc = Document()
    t = doc.add_heading("PhÃ¢n tÃ­ch dá»± Ã¡n LocalNavBot", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "TÃ i liá»‡u mÃ´ táº£ kiáº¿n trÃºc, luá»“ng xá»­ lÃ½, giáº£i thÃ­ch cÃ¡c khá»‘i mÃ£ chÃ­nh, "
        "Ä‘Ã¡nh giÃ¡ Ä‘Ãºng/sai vÃ  Æ°u/nhÆ°á»£c Ä‘iá»ƒm. "
        "NgÃ y táº¡o: tham chiáº¿u mÃ£ nguá»“n trong thÆ° má»¥c dá»± Ã¡n."
    )

    # --- Tá»•ng quan ---
    add_heading(doc, "1. Tá»•ng quan há»‡ thá»‘ng", 1)
    add_para(
        doc,
        "LocalNavBot lÃ  á»©ng dá»¥ng Ä‘iá»u hÆ°á»›ng Ä‘á»‹a phÆ°Æ¡ng káº¿t há»£p: API FastAPI, "
        "bá»™ Ä‘á»‹nh tuyáº¿n (Valhalla + fallback OSMnx/A*), cÆ¡ sá»Ÿ dá»¯ liá»‡u SQLite, "
        "VPR (Visual Place Recognition) Ä‘á»ƒ nháº­n diá»‡n Ä‘á»‹a Ä‘iá»ƒm tá»« áº£nh, "
        "vÃ  bot há»™i thoáº¡i dÃ¹ng LLM (Anthropic/OpenAI/Ollama).",
    )
    add_heading(doc, "1.1. CÃ¡c tÃ­nh nÄƒng chÃ­nh (theo mÃ£ nguá»“n)", 2)
    add_bullets(
        doc,
        [
            "Chat / streaming / WebSocket: há»i Ä‘Ã¡p Ä‘iá»u hÆ°á»›ng, ngá»¯ cáº£nh GPS.",
            "TÃ¬m tuyáº¿n: giáº£i mÃ£ Ä‘á»‹a chá»‰ (DB local + Nominatim náº¿u báº­t), gá»i Valhalla hoáº·c A* trÃªn Ä‘á»“ thá»‹ OSM.",
            "Dá»¯ liá»‡u local: locations, images, POI, cáº¡nh tÃ¹y chá»‰nh (háº»m/Ä‘Æ°á»ng táº¯t), quan sÃ¡t giao thÃ´ng/mÃ´i trÆ°á»ng.",
            "VPR: truy váº¥n áº£nh, index/rebuild FAISS.",
            "Báº£n Ä‘á»“ Folium, isochrone, heatmap táº¯c ngháº½n, upload áº£nh batch.",
            "CLI (Typer): serve, index, demo, status.",
        ],
    )

    # --- Kiáº¿n trÃºc ---
    add_heading(doc, "2. Kiáº¿n trÃºc thÆ° má»¥c", 1)
    add_bullets(
        doc,
        [
            "main.py â€” Ä‘iá»ƒm vÃ o nhanh, gá»i Typer app trong scripts/cli.py.",
            "config/settings.py â€” cáº¥u hÃ¬nh Pydantic (.env).",
            "core/ â€” database, VPR, OCR, YOLO landmark, street view, traffic, environment.",
            "routing/ â€” NavRouter, TrafficHeuristic, OSMGraph, ValhallaClient, render.",
            "bot/ â€” NavBot (LLM + intent), session_manager.",
            "web/ â€” FastAPI app, routes modular, ui.html.",
            "scripts/cli.py â€” lá»‡nh dÃ²ng lá»‡nh.",
        ],
    )

    # --- main.py ---
    add_heading(doc, "3. PhÃ¢n tÃ­ch tá»«ng file quan trá»ng", 1)

    add_heading(doc, "3.1. main.py", 2)
    add_para(doc, "Chá»©c nÄƒng:", bold=True)
    add_bullets(
        doc,
        [
            "ChÃ¨n thÆ° má»¥c gá»‘c project vÃ o sys.path Ä‘á»ƒ import á»•n Ä‘á»‹nh khi cháº¡y trá»±c tiáº¿p.",
            "Gá»i app() cá»§a Typer tá»« scripts.cli.",
        ],
    )
    add_para(doc, "ÄÃ¡nh giÃ¡:", bold=True)
    add_bullets(
        doc,
        [
            "ÄÃºng: pattern phá»• biáº¿n cho package chÆ°a cÃ i pip editable.",
            "Æ¯u: tá»‘i giáº£n, dá»… Ä‘á»c.",
            "NhÆ°á»£c: trÃ¹ng logic insert path vá»›i cli.py (hai nÆ¡i).",
        ],
    )

    # --- settings ---
    add_heading(doc, "3.2. config/settings.py", 2)
    add_para(doc, "Chá»©c nÄƒng:", bold=True)
    add_bullets(
        doc,
        [
            "BASE_DIR vÃ  class Settings káº¿ thá»«a BaseSettings: Ä‘á»c .env, Ä‘Æ°á»ng dáº«n data, Valhalla, OSM area, trá»ng sá»‘ routing, LLM, CORS.",
            "setup_dirs(): táº¡o thÆ° má»¥c dá»¯ liá»‡u.",
            "Property cors_origin_list, ocr_language_list: parse chuá»—i CSV.",
        ],
    )
    add_para(doc, "ÄÃ¡nh giÃ¡:", bold=True)
    add_bullets(
        doc,
        [
            "ÄÃºng: tÃ¡ch cáº¥u hÃ¬nh rÃµ, type hints Literal cho llm_provider.",
            "Æ¯u: pydantic-settings chuáº©n, extra=ignore an toÃ n khi .env thá»«a key.",
            "NhÆ°á»£c: máº·c Ä‘á»‹nh device=cuda cÃ³ thá»ƒ gÃ¢y lá»—i trÃªn mÃ¡y khÃ´ng cÃ³ GPU náº¿u khÃ´ng Ä‘á»•i .env.",
        ],
    )

    # --- database ---
    add_heading(doc, "3.3. core/database.py", 2)
    add_para(doc, "Chá»©c nÄƒng:", bold=True)
    add_bullets(
        doc,
        [
            "SCHEMA_SQL: báº£ng locations, images, pois, custom_edges, traffic_observations, environmental_observations, nav_sessions; index lat/lon.",
            "Class Database: aiosqlite, init(), CRUD vÃ  truy váº¥n lÃ¢n cáº­n (nearby), search.",
        ],
    )
    add_para(doc, "ÄÃ¡nh giÃ¡:", bold=True)
    add_bullets(
        doc,
        [
            "ÄÃºng: WAL, foreign_keys; phÃ¹ há»£p prototype/local.",
            "Æ¯u: async, schema Ä‘áº§y Ä‘á»§ cho use case Ä‘iá»u hÆ°á»›ng + crowdsourcing.",
            "NhÆ°á»£c: khÃ´ng dÃ¹ng SpatiaLite Ä‘áº§y Ä‘á»§ trong DDL (comment nÃ³i spatial); truy váº¥n gáº§n dá»±a bbox/Ä‘á»™ Ä‘Æ¡n giáº£n tÃ¹y implementation.",
        ],
    )

    # --- router ---
    add_heading(doc, "3.4. routing/router.py", 2)
    add_para(doc, "Chá»©c nÄƒng (theo khá»‘i mÃ£):", bold=True)
    add_bullets(
        doc,
        [
            "RouteStep / Route: bÆ°á»›c chá»‰ dáº«n, geometry, phÃ¢n tÃ­ch phá»¥.",
            "TrafficHeuristic: warm_cache theo giá», congestion_factor káº¿t há»£p DB + peak_hours, edge_weight nhÃ¢n trá»ng sá»‘ + environmental_analyzer.",
            "OSMGraph: load/cache graphml hoáº·c táº£i osmnx, patch_custom_edges ná»‘i cáº¡nh DB vÃ o Ä‘á»“ thá»‹, add_travel_times.",
            "ValhallaClient: GET /status, POST /route vá»›i date_time vÃ  costing_options.",
            "_vn_instruction: template tiáº¿ng Viá»‡t cho maneuver.",
            "OSMNXRouter / SmartOSMNXRouter: A* vá»›i _weight, nhiá»u profile, route_score Ä‘á»ƒ chá»n tuyáº¿n.",
            "NavRouter.init: load OSM, patch edges, warm heuristic, check Valhalla; find_route Æ°u tiÃªn Valhalla rá»“i fallback.",
            "resolve_location: DB locations/pois trÆ°á»›c, sau Ä‘Ã³ Nominatim náº¿u allow_remote_geocoding.",
            "HÃ m há»— trá»£: haversine, bearing, decode polyline Valhalla.",
        ],
    )
    add_para(doc, "ÄÃ¡nh giÃ¡:", bold=True)
    add_bullets(
        doc,
        [
            "ÄÃºng: kiáº¿n trÃºc 2 táº§ng Valhalla + fallback máº¡nh khi Docker Valhalla táº¯t.",
            "Æ¯u: heuristic giá» cao Ä‘iá»ƒm + quan sÃ¡t DB + mÃ´i trÆ°á»ng; Ä‘a profile rerank cÃ³ Ã½ nghÄ©a.",
            "Sai / lá»—i cáº§n sá»­a: trong SmartOSMNXRouter._route_from_path, chuá»—i tiáº¿ng Viá»‡t cho street fallback vÃ  bÆ°á»›c arrive bá»‹ lá»—i encoding "
            '(vÃ­ dá»¥ hiá»ƒn thá»‹ "Ã„\'Ã†Â°Ã¡Â»\u009dng" thay vÃ¬ "Ä‘Æ°á»ng") â€” file .py cÃ³ thá»ƒ Ä‘Ã£ lÆ°u sai encoding hoáº·c copy-paste lá»—i; nÃªn thá»‘ng nháº¥t UTF-8 vÃ  dÃ¹ng cÃ¹ng literal vá»›i _vn_instruction.',
            "NhÆ°á»£c: heuristic A* dÃ¹ng lat/lon node u cho má»i edge (u,v) â€” gáº§n Ä‘Ãºng; Valhalla parse geometry chá»‰ láº¥y Ä‘iá»ƒm theo maneuver, cÃ³ thá»ƒ thÆ°a.",
            "LÆ°u Ã½: total_duration_s trong OSM path Ä‘ang cá»™ng _weight (Ä‘Ã£ nhÃ¢n há»‡ sá»‘) nhÆ°ng RouteStep.duration_s váº«n lÃ  travel_time gá»‘c â€” cÃ³ thá»ƒ lá»‡ch giá»¯a tá»•ng thá»i gian bÆ°á»›c vÃ  tá»•ng weighted.",
        ],
    )

    # --- nav_bot ---
    add_heading(doc, "3.5. bot/nav_bot.py", 2)
    add_para(doc, "Chá»©c nÄƒng:", bold=True)
    add_bullets(
        doc,
        [
            "SYSTEM_PROMPT: persona LocalNavBot tiáº¿ng Viá»‡t.",
            "NavigationIntent: dataclass intent tá»« LLM.",
            "_img_to_base64, _resize_for_api: chuáº©n bá»‹ áº£nh JPEG base64 cho API.",
            "LLMClient: Anthropic vs OpenAI/Ollama, chat sync/stream, parse_intent (JSON tá»« LLM).",
            "_format_route_for_bot: markdown tuyáº¿n + cáº£nh bÃ¡o peak + áº£nh bÆ°á»›c.",
            "_attach_images_to_steps: stub (pass) â€” chÆ°a dÃ¹ng.",
            "NavBot.ask: context GPS â†’ parse_intent â†’ _fallback_intent â†’ route/find_poi/identify â†’ ghÃ©p system context â†’ history â†’ LLM.",
            "_handle_route, _handle_identify, _build_visual_context (YOLO/OCR/VPR).",
            "add_location_with_images, rebuild_vpr_index.",
        ],
    )
    add_para(doc, "ÄÃ¡nh giÃ¡:", bold=True)
    add_bullets(
        doc,
        [
            "ÄÃºng: tÃ¡ch provider LLM; fallback intent theo tá»« khÃ³a tiáº¿ng Viá»‡t khÃ´ng dáº¥u há»¯u Ã­ch.",
            "Æ¯u: luá»“ng rÃµ: intent â†’ dá»¯ liá»‡u â†’ LLM tráº£ lá»i cÃ³ context.",
            "NhÆ°á»£c / rá»§i ro: parse_intent dÃ¹ng NavigationIntent(**data) â€” náº¿u LLM tráº£ thá»«a key sáº½ lá»—i TypeError (catch rÆ¡i vá» chat).",
            "NhÆ°á»£c: strip markdown ```json``` cÃ³ thá»ƒ khÃ´ng sáº¡ch náº¿u cÃ³ text trÆ°á»›c/sau JSON.",
            "Anthropic message vá»›i image: trong ask(), user_content dÃ¹ng key \"image\" â€” cáº§n khá»›p Ä‘Ãºng schema SDK Anthropic (náº¿u khÃ´ng khá»›p sáº½ lá»—i runtime); OpenAI path dÃ¹ng messages Ä‘Æ¡n giáº£n hÆ¡n.",
            "Streaming: khi stream=True, history assistant khÃ´ng Ä‘Æ°á»£c append trong ask() â€” Ä‘Ãºng Ã½ trÃ¡nh partial; caller (web) pháº£i tá»± lÆ°u full.",
            "HÃ m _attach_images_to_steps dead code â€” nÃªn xÃ³a hoáº·c implement.",
        ],
    )

    # --- app.py ---
    add_heading(doc, "3.6. web/app.py", 2)
    add_para(doc, "Chá»©c nÄƒng:", bold=True)
    add_bullets(
        doc,
        [
            "FastAPI app, CORS, include_router nhiá»u module.",
            "startup: db.init, NavRouter.init, VPREngine (try/except), NavBot, traffic/environment refresh, session_manager.start.",
            "Legacy REST: chat, chat stream SSE, websocket chat+GPS, route, GPS, traffic, upload, map HTML, VPR, jobs, status.",
            "GET / phá»¥c vá»¥ ui.html.",
        ],
    )
    add_para(doc, "ÄÃ¡nh giÃ¡:", bold=True)
    add_bullets(
        doc,
        [
            "ÄÃºng: tÃ¡ch router; 503 khi router/bot chÆ°a ready.",
            "Æ¯u: API legacy Ä‘áº§y Ä‘á»§ cho UI vÃ  tÃ­ch há»£p.",
            "NhÆ°á»£c: legacy_chat táº¡o NavBot má»›i má»—i request nhÆ°ng copy history tá»« session â€” khÃ´ng dÃ¹ng singleton _bot (cÃ³ thá»ƒ cá»‘ Ã½ Ä‘á»ƒ trÃ¡nh shared state); cáº§n Ä‘á»“ng bá»™ vá»›i _vpr.",
            "FastAPI on_event('startup') deprecated trong báº£n má»›i â€” cÃ³ thá»ƒ migrate lifespan.",
        ],
    )

    # --- cli ---
    add_heading(doc, "3.7. scripts/cli.py", 2)
    add_para(doc, "Chá»©c nÄƒng: lá»‡nh serve (uvicorn), index áº£nh + VPR, add-location, demo, status.", bold=True)
    add_para(doc, "ÄÃ¡nh giÃ¡: UX CLI tá»‘t (Rich/Typer); async subprocess trong cÃ¡c lá»‡nh dÃ i.", bold=True)

    # --- Báº£ng luá»“ng ---
    add_heading(doc, "4. Luá»“ng xá»­ lÃ½ tiÃªu biá»ƒu", 1)
    add_para(doc, "4.1. NgÆ°á»i dÃ¹ng há»i \"Äi tá»« A Ä‘áº¿n B\"", bold=True)
    add_bullets(
        doc,
        [
            "LLM parse_intent â†’ intent route + origin/destination (hoáº·c fallback tá»« text).",
            "resolve_location cho Ä‘iá»ƒm Ä‘áº§u/cuá»‘i.",
            "NavRouter.find_route (Valhalla hoáº·c OSMnx).",
            "_attach_images_to_route láº¥y áº£nh DB + street view.",
            "_format_route_for_bot Ä‘Æ°a vÃ o system context; LLM sinh cÃ¢u tráº£ lá»i cuá»‘i.",
        ],
    )
    add_para(doc, "4.2. Upload áº£nh cÃ³ GPS", bold=True)
    add_bullets(
        doc,
        [
            "validate_upload, lÆ°u file, read_gps_exif, add_location/add_image, optional VPR index_image.",
        ],
    )

    # --- Æ¯u nhÆ°á»£c tá»•ng ---
    add_heading(doc, "5. Æ¯u Ä‘iá»ƒm vÃ  háº¡n cháº¿ tá»•ng thá»ƒ", 1)
    add_heading(doc, "5.1. Æ¯u Ä‘iá»ƒm", 2)
    add_bullets(
        doc,
        [
            "BÃ¡m sÃ¡t bÃ i toÃ¡n Ä‘á»‹a phÆ°Æ¡ng: POI/cáº¡nh tÃ¹y chá»‰nh, áº£nh minh há»a, VPR.",
            "Routing linh hoáº¡t: engine chuyÃªn nghiá»‡p + offline fallback.",
            "Má»Ÿ rá»™ng module web/routes, core tÃ¡ch báº¡ch.",
            "Observability cÆ¡ báº£n: loguru, status endpoint.",
        ],
    )
    add_heading(doc, "5.2. Háº¡n cháº¿ / rá»§i ro", 2)
    add_bullets(
        doc,
        [
            "Phá»¥ thuá»™c nhiá»u service (Valhalla, LLM API, GPU cho VPR/YOLO).",
            "Äá»™ phá»©c táº¡p triá»ƒn khai: torch, faiss-gpu theo ná»n táº£ng.",
            "Má»™t sá»‘ chuá»—i Unicode lá»—i trong router (cáº§n rÃ  soÃ¡t encoding).",
            "Báº£o máº­t: CORS máº·c Ä‘á»‹nh localhost; cáº§n harden khi public.",
            ".env cÃ³ thá»ƒ chá»©a secret â€” khÃ´ng commit (Ä‘Ã£ cÃ³ .env.example).",
        ],
    )

    add_heading(doc, "6. Gá»£i Ã½ cáº£i tiáº¿n ngáº¯n", 1)
    add_bullets(
        doc,
        [
            "Sá»­a literal tiáº¿ng Viá»‡t bá»‹ corrupt trong SmartOSMNXRouter.",
            "parse_intent: dÃ¹ng model_validate vá»›i dict loáº¡i bá» key thá»«a hoáº·c json.loads + chá»‰ láº¥y field há»£p lá»‡.",
            "Gá»¡ hoáº·c hoÃ n thiá»‡n _attach_images_to_steps.",
            "ThÃªm test tÃ­ch há»£p nhá» cho resolve_location vÃ  find_route mock Valhalla.",
            "Lifespan FastAPI thay startup event.",
        ],
    )

    doc.add_paragraph("")
    add_para(doc, "â€” Háº¿t tÃ i liá»‡u â€”", bold=True)

    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    main()
